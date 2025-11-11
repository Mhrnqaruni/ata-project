# /ata-backend/app/services/page_count_service.py

"""
This service provides functionality to count the total number of pages
across multiple uploaded files (PDF, DOCX, and images).

For PDFs: Counts actual pages in the document.
For DOCX: Counts pages (approximation based on page breaks and content).
For Images: Each image counts as 1 page.
"""

import fitz  # PyMuPDF
from docx import Document
from fastapi import UploadFile
from typing import List
from io import BytesIO


async def count_pages_in_file(file: UploadFile) -> int:
    """
    Count the number of pages in a single file based on its type.

    Args:
        file: The uploaded file (PDF, DOCX, or image)

    Returns:
        int: Number of pages in the file

    Raises:
        ValueError: If the file type is not supported
    """
    file_content = await file.read()
    file_type = file.content_type.lower() if file.content_type else ""
    filename_lower = file.filename.lower() if file.filename else ""

    # Reset file pointer for potential reuse
    await file.seek(0)

    # Handle PDF files
    if "pdf" in file_type or filename_lower.endswith(".pdf"):
        try:
            pdf_document = fitz.open(stream=file_content, filetype="pdf")
            page_count = pdf_document.page_count
            pdf_document.close()
            return page_count
        except Exception as e:
            raise ValueError(f"Error reading PDF file '{file.filename}': {str(e)}")

    # Handle DOCX files
    elif "wordprocessingml" in file_type or filename_lower.endswith(".docx"):
        try:
            doc = Document(BytesIO(file_content))
            # Count pages based on explicit page breaks
            page_count = 1  # Start with 1 page
            for paragraph in doc.paragraphs:
                # Check if paragraph contains a page break
                if '\f' in paragraph.text or '\x0c' in paragraph.text:
                    page_count += 1
                # Check for page break in runs
                for run in paragraph.runs:
                    if run._element.xml.find('w:br') != -1 and 'w:type="page"' in run._element.xml:
                        page_count += 1

            # If no explicit page breaks found, estimate based on content
            if page_count == 1 and len(doc.paragraphs) > 0:
                # Rough estimation: ~45 lines per page, ~2 paragraphs per line on average
                estimated_lines = sum(1 for p in doc.paragraphs if p.text.strip())
                page_count = max(1, estimated_lines // 45)

            return page_count
        except Exception as e:
            raise ValueError(f"Error reading DOCX file '{file.filename}': {str(e)}")

    # Handle image files (JPEG, PNG, etc.)
    elif any(img_type in file_type for img_type in ["image/jpeg", "image/png", "image/jpg"]) or \
         any(filename_lower.endswith(ext) for ext in [".jpg", ".jpeg", ".png"]):
        # Each image counts as 1 page
        return 1

    else:
        raise ValueError(f"Unsupported file type: {file.filename} ({file_type})")


async def count_total_pages(files: List[UploadFile]) -> int:
    """
    Count the total number of pages across all uploaded files.

    Args:
        files: List of uploaded files (can be PDFs, DOCX, or images)

    Returns:
        int: Total number of pages across all files
    """
    total_pages = 0

    for file in files:
        try:
            pages = await count_pages_in_file(file)
            total_pages += pages
        except ValueError as e:
            # Log the error but continue processing other files
            print(f"Warning: {str(e)}")
            # Assume 1 page for files that can't be processed
            total_pages += 1

    return total_pages
