# /app/services/class_helpers/file_processors.py

import io
from typing import List, Dict
import pandas as pd
import fitz
from PIL import Image

try:
    import docx
except ImportError:
    docx = None

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Robustly extracts text from a .docx file, including paragraphs and tables.
    """
    if docx is None:
        raise ImportError("The 'python-docx' library is not installed. Please run 'pip install python-docx'.")
    stream = io.BytesIO(file_bytes)
    document = docx.Document(stream)
    full_text = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return "\n".join(full_text)

def extract_students_from_tabular(file_bytes: bytes, is_excel: bool) -> List[Dict]:
    """
    Directly parses student data from Excel or CSV files, bypassing AI.
    """
    stream = io.BytesIO(file_bytes)
    try:
        df = pd.read_excel(stream) if is_excel else pd.read_csv(stream)
    except Exception as e:
        raise ValueError(f"Could not parse the provided tabular file: {e}")
    name_col = next((col for col in df.columns if 'name' in col.lower()), None)
    id_col = next((col for col in df.columns if 'id' in col.lower()), None)
    if not name_col or not id_col:
        raise ValueError("Tabular file must contain columns with 'name' and 'id' in their titles.")
    df = df.rename(columns={name_col: 'name', id_col: 'studentId'})
    df['studentId'] = df['studentId'].astype(str)
    return df[['name', 'studentId']].to_dict('records')

def convert_pdf_first_page_to_png_bytes(pdf_bytes: bytes) -> bytes:
    """
    Converts the first page of a PDF file into PNG image bytes for multi-modal AI calls.
    """
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if not len(doc): raise ValueError("PDF file is empty.")
        page = doc.load_page(0)
        pix = page.get_pixmap(dpi=200)
        doc.close()
        img_stream = io.BytesIO()
        Image.frombytes("RGB", [pix.width, pix.height], pix.samples).save(img_stream, format="PNG")
        return img_stream.getvalue()
    except Exception as e:
        raise ValueError(f"Could not process the provided PDF file as an image: {e}")