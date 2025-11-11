# /ata-backend/app/services/assessment_service.py (REVISED AND CORRECTED)

import uuid, json, os, asyncio, shutil
from fastapi import UploadFile, Depends
from typing import List, Dict, Optional, Union, Tuple
from decimal import Decimal, InvalidOperation
from datetime import datetime, timezone
from io import BytesIO
from docx import Document

# Import all services and helpers
from . import ocr_service, gemini_service, prompt_library, report_service, zip_service, library_service, pdf_service
from .database_service import DatabaseService, get_db_service
from ..models import assessment_model
from ..db.models.assessment_models import ResultStatus, FinalizedBy
from .assessment_helpers import job_creation, grading_pipeline, data_assembly, document_parser, analytics_and_matching

ASSESSMENT_UPLOADS_DIR = "assessment_uploads"

def finalize_question(
    three_grades: List[Optional[Decimal]],
    three_comments: List[Optional[str]],
    max_score: float
) -> Dict:
    """
    Computes consensus for a single question from three AI model runs.
    This version is more robust and correctly handles edge cases.
    """
    valid_runs = [(g, c) for g, c in zip(three_grades, three_comments) if g is not None]

    if len(valid_runs) < 2:
        return {"status": ResultStatus.PENDING_REVIEW, "grade": None, "feedback": None, "finalized_by": None}

    grade_groups: Dict[int, list] = {}
    for i in range(len(valid_runs)):
        g1, c1 = valid_runs[i]
        matched = False
        for group_id, group in grade_groups.items():
            if any(abs(g1 - g_existing) <= Decimal("0.1") for g_existing, _ in group):
                grade_groups[group_id].append((g1, c1))
                matched = True
                break
        if not matched:
            grade_groups[len(grade_groups)] = [(g1, c1)]

    for group in grade_groups.values():
        if len(group) >= 2:
            # Correctly average the grades of the majority group
            total_grade = sum(g for g, c in group)
            avg_grade = total_grade / len(group)

            # Use the comment from the first member of the majority group
            final_comment = group[0][1]

            return {
                "status": ResultStatus.AI_GRADED,
                "grade": float(round(avg_grade, 2)),
                "feedback": final_comment,
                "finalized_by": FinalizedBy.AI,
            }

    return {"status": ResultStatus.PENDING_REVIEW, "grade": None, "feedback": None, "finalized_by": None}


class AssessmentService:
    def __init__(self, db: DatabaseService = Depends(get_db_service)):
        self.db = db

    async def parse_document_for_review(self, question_file: UploadFile, answer_key_file: Optional[UploadFile], class_id: str, assessment_name: str) -> Dict:
        return await document_parser.parse_document_to_config(question_file, answer_key_file, class_id, assessment_name)

    async def create_new_assessment_job_v2(
        self, config: assessment_model.AssessmentConfigV2, answer_sheets: List[UploadFile], user_id: str
    ) -> Dict:
        job_id = f"job_{uuid.uuid4().hex[:16]}"

        # Count total pages across all student submissions
        from . import page_count_service
        print(f"[ASSESSMENT SERVICE] Counting pages for {len(answer_sheets)} files...")
        total_pages = await page_count_service.count_total_pages(answer_sheets)
        print(f"[ASSESSMENT SERVICE] Total pages counted: {total_pages}")

        answer_sheet_data = job_creation._save_uploaded_files(job_id, answer_sheets)
        job_creation._create_initial_job_records_v2(self.db, job_id, config, answer_sheet_data, user_id, total_pages)

        response = {
            "jobId": job_id,
            "status": assessment_model.JobStatus.QUEUED.value,
            "message": "Assessment job created.",
            "totalPages": total_pages
        }
        print(f"[ASSESSMENT SERVICE] Returning response: {response}")
        return response

    async def create_job_with_manual_uploads(
        self, config: assessment_model.AssessmentConfigV2, form_data, outsider_names: List[Dict[str, str]], user_id: str
    ) -> Dict:
        """
        Creates a job, processes manually uploaded files keyed by student/outsider ID,
        and links everything in the database in a single transaction.
        """
        print(f"[MANUAL-UPLOAD-START] ========== Starting Manual Upload Processing ==========")
        print(f"[MANUAL-UPLOAD-START] Assessment: {config.assessmentName}")
        print(f"[MANUAL-UPLOAD-START] Outsider names map: {outsider_names}")

        # 1. Create the base job and questions, but with no answer sheets yet.
        job_id = f"job_{uuid.uuid4().hex[:16]}"
        job_creation._create_initial_job_records_v2(self.db, job_id, config, [], user_id, is_manual_upload=True)
        print(f"[MANUAL-UPLOAD-START] Created job: {job_id}")

        # Create a map for quick lookup of outsider names by their temporary frontend ID
        outsider_name_map = {item['id']: item['name'] for item in outsider_names}

        # 2. Group files by their key (to handle multiple files per student)
        # DIAGNOSTIC: First, let's see ALL items in FormData
        all_items = []
        for key, value in form_data.multi_items():
            if hasattr(value, 'filename'):
                all_items.append((key, value.filename, value.content_type))
            else:
                all_items.append((key, str(value)[:50]))
        print(f"[MANUAL-UPLOAD-FORMDATA] Raw FormData items: {all_items}")

        files_by_key = {}
        for key, file in form_data.multi_items():
            if not key.endswith('_files'):
                continue
            if key not in files_by_key:
                files_by_key[key] = []
            files_by_key[key].append(file)

        print(f"[MANUAL-UPLOAD-GROUPED] Files grouped by key: {[(k, len(v)) for k, v in files_by_key.items()]}")

        # 3. Process each set of files
        for key, files in files_by_key.items():
            print(f"[DEBUG] Processing key '{key}' with {len(files)} files")
            # Skip if no files were actually uploaded for this key
            if not files or not files[0].filename:
                continue

            parts = key.replace('_files', '').split('_')
            entity_type = parts[0] # 'student' or 'outsider'
            entity_id = '_'.join(parts[1:]) # Re-join in case the ID had underscores

            # 3. Compress, merge, and save the PDF for the entity
            try:
                compressed_images = [pdf_service.compress_image(await f.read()) for f in files]
                pdf_bytes = pdf_service.merge_images_to_pdf(compressed_images)

                job_dir = os.path.join(ASSESSMENT_UPLOADS_DIR, job_id)
                os.makedirs(job_dir, exist_ok=True)

                pdf_filename = f"manual_{entity_type}_{entity_id}.pdf"
                pdf_path = os.path.join(job_dir, pdf_filename)
                with open(pdf_path, "wb") as f:
                    f.write(pdf_bytes)

                file_info = {"path": pdf_path, "contentType": "application/pdf"}

                # 4. Create DB records (Result, AnswerSheet, and possibly OutsiderStudent)
                if entity_type == 'student':
                    print(f"[DEBUG] Creating results for student ID: {entity_id}")
                    analytics_and_matching._create_results_for_entity(self.db, job_id, entity_id, 'student', config, file_info, user_id)
                elif entity_type == 'outsider':
                    outsider_name = outsider_name_map.get(entity_id, "Unknown Outsider")
                    print(f"[DEBUG] Creating outsider with name '{outsider_name}' for entity_id '{entity_id}'")
                    new_outsider = self.db.add_outsider_student({
                        "name": outsider_name,
                        "assessment_id": job_id
                    })
                    print(f"[DEBUG] Outsider created with DB ID: {new_outsider.id}")
                    analytics_and_matching._create_results_for_entity(self.db, job_id, new_outsider.id, 'outsider', config, file_info, user_id)

            except Exception as e:
                print(f"Error processing files for key {key}: {e}")
                # Optionally, decide if a single failure should halt the whole process
                # For now, we'll log and continue.
                continue

        return { "jobId": job_id, "status": assessment_model.JobStatus.QUEUED.value, "message": "Assessment job with manual uploads created." }

    async def process_manual_submission(self, job_id: str, user_id: str, images: List[UploadFile], student_id: Optional[str] = None, outsider_name: Optional[str] = None) -> Dict:
        if not student_id and not outsider_name:
            raise ValueError("Either a student_id or an outsider_name must be provided.")

        print(f"[MANUAL-SUBMIT] Received {len(images)} images for job={job_id}, student_id={student_id}, outsider_name={outsider_name}")

        # 1. Compress and merge images into a single PDF
        compressed_images = [pdf_service.compress_image(await img.read()) for img in images]
        pdf_bytes = pdf_service.merge_images_to_pdf(compressed_images)

        # 2. Get the assessment config
        job = self.db.get_assessment_job(job_id, user_id)
        if not job:
            raise ValueError(f"Job {job_id} not found.")
        config = analytics_and_matching.normalize_config_to_v2(job)

        # 3. Handle student vs outsider
        if student_id:
            # It's a rostered student
            print(f"[MANUAL-SUBMIT] Processing rostered student: {student_id}")

            # Save PDF
            job_dir = os.path.join(ASSESSMENT_UPLOADS_DIR, job_id)
            os.makedirs(job_dir, exist_ok=True)
            pdf_filename = f"manual_{student_id}.pdf"
            pdf_path = os.path.join(job_dir, pdf_filename)

            with open(pdf_path, "wb") as f:
                f.write(pdf_bytes)

            file_info = {"path": pdf_path, "contentType": "application/pdf"}
            analytics_and_matching._create_results_for_entity(self.db, job_id, student_id, 'student', config, file_info, user_id)
            return {"message": f"Submission for student {student_id} processed successfully."}
        else:
            # It's an outsider student - check if they already exist for this job
            print(f"[MANUAL-SUBMIT] Processing outsider: {outsider_name}")
            existing_outsider = self.db.get_outsider_by_name_and_job(outsider_name, job_id, user_id)

            if existing_outsider:
                # Outsider already exists - append new pages to their existing PDF
                print(f"[MANUAL-SUBMIT] Outsider '{outsider_name}' already exists (ID: {existing_outsider.id}), appending pages")

                # Get existing PDF path
                job_dir = os.path.join(ASSESSMENT_UPLOADS_DIR, job_id)
                existing_pdf_filename = f"manual_outsider_{existing_outsider.id}.pdf"
                existing_pdf_path = os.path.join(job_dir, existing_pdf_filename)

                # Merge new PDF with existing PDF
                if os.path.exists(existing_pdf_path):
                    with open(existing_pdf_path, "rb") as f:
                        existing_pdf_bytes = f.read()
                    # Merge existing PDF + new PDF
                    merged_pdf_bytes = pdf_service.merge_pdfs([existing_pdf_bytes, pdf_bytes])
                else:
                    # Existing PDF not found, use new PDF only
                    merged_pdf_bytes = pdf_bytes

                # Save merged PDF
                with open(existing_pdf_path, "wb") as f:
                    f.write(merged_pdf_bytes)

                return {"message": f"Submission for outsider {outsider_name} appended successfully."}
            else:
                # Create new outsider
                print(f"[MANUAL-SUBMIT] Creating new outsider: {outsider_name}")
                new_outsider = self.db.add_outsider_student({
                    "name": outsider_name,
                    "assessment_id": job_id
                })

                # Save PDF
                job_dir = os.path.join(ASSESSMENT_UPLOADS_DIR, job_id)
                os.makedirs(job_dir, exist_ok=True)
                pdf_filename = f"manual_outsider_{new_outsider.id}.pdf"
                pdf_path = os.path.join(job_dir, pdf_filename)

                with open(pdf_path, "wb") as f:
                    f.write(pdf_bytes)

                file_info = {"path": pdf_path, "contentType": "application/pdf"}
                analytics_and_matching._create_results_for_entity(self.db, job_id, new_outsider.id, 'outsider', config, file_info, user_id)
                return {"message": f"Submission for outsider {outsider_name} processed successfully."}

    def create_new_assessment_job(
        self, config: assessment_model.AssessmentConfig, answer_sheets: List[UploadFile], user_id: str
    ) -> Dict:
        job_id = f"job_{uuid.uuid4().hex[:16]}"
        answer_sheet_data = job_creation._save_uploaded_files(job_id, answer_sheets)
        job_creation._create_initial_job_records(self.db, job_id, config, answer_sheet_data, user_id)
        return { "jobId": job_id, "status": assessment_model.JobStatus.QUEUED.value, "message": "Assessment job created." }

    def delete_assessment_job(self, job_id: str, user_id: str) -> bool:
        was_deleted = self.db.delete_assessment_job(job_id=job_id, user_id=user_id)
        if not was_deleted:
            return False
        job_dir = os.path.join(ASSESSMENT_UPLOADS_DIR, job_id)
        if os.path.isdir(job_dir):
            shutil.rmtree(job_dir)
        return True

    async def _invoke_single_grading_run_vision(
        self, prompt: str, file_bytes: bytes, mime_type: str, job_id: str, entity_id: str, is_outsider: bool, run_index: int, user_id: str
    ) -> Dict:
        print(f"[RUN-START-VISION] job={job_id} entity={entity_id} is_outsider={is_outsider} run={run_index} model=gemini-2.5-flash file_size={len(file_bytes)} prompt_chars={len(prompt)}")

        all_results_for_job = self.db.get_all_results_for_job(job_id, user_id)
        if is_outsider:
            entity_results = [r for r in all_results_for_job if r.outsider_student_id == entity_id]
        else:
            entity_results = [r for r in all_results_for_job if r.student_id == entity_id]
        all_question_ids = [r.question_id for r in entity_results]

        # Get student name for logging
        student_name = "Unknown"
        if not is_outsider:
            student_obj = self.db.get_student_by_id(entity_id, user_id)
            if student_obj:
                student_name = student_obj.name
        else:
            outsider_obj = self.db.get_outsider_student_by_id(entity_id, user_id)
            if outsider_obj:
                student_name = outsider_obj.name

        try:
            # Use vision-based processing with JSON mode
            result = await gemini_service.process_file_with_vision_json(
                file_bytes=file_bytes,
                mime_type=mime_type,
                prompt=prompt,
                temperature=0.1,
                log_context=f"GRADE-STUDENT [{student_name}] Run-{run_index}"
            )
            parsed_results = result['data']
            tokens_used = result['tokens']
            print(f"[RUN-VISION-PARSED] job={job_id} entity={entity_id} run={run_index} results_count={len(parsed_results.get('results', []))}")

            # Prepare the base payload for the AI model run log
            base_payload = {
                "job_id": job_id,
                "run_index": run_index,
                "student_id": None,
                "outsider_student_id": None
            }
            if is_outsider:
                base_payload["outsider_student_id"] = entity_id
            else:
                base_payload["student_id"] = entity_id

            for result_data in parsed_results.get('results', []):
                question_id = result_data.get('question_id')
                if not question_id or question_id not in all_question_ids:
                    continue

                grade = grading_pipeline._safe_float_convert(result_data.get('grade'))
                comment = result_data.get('feedback', '')
                extracted_answer = result_data.get('extractedAnswer', '')

                print(f"[RUN-PARSED-VISION] job={job_id} entity={entity_id} run={run_index} q={question_id} grade={grade} hasExtracted={bool(extracted_answer)}")

                run_payload = {
                    **base_payload,
                    "question_id": question_id,
                    "raw_json": result_data,
                    "grade": grade,
                    "comment": comment
                }
                self.db.create_ai_model_run(**run_payload)
                print(f"[RUN-SAVE-VISION] job={job_id} entity={entity_id} run={run_index} q={question_id} grade={grade} commentChars={len(comment or '')}")

                if run_index == 0:
                    self.db.update_result_extracted_answer(job_id, entity_id, is_outsider, question_id, extracted_answer, user_id)
            return {'success': True, 'tokens': tokens_used}
        except Exception as e:
            print(f"Error in vision-based grading run {run_index} for entity {entity_id} in job {job_id}: {e}")

            error_payload = {
                "job_id": job_id,
                "run_index": run_index,
                "raw_json": {"error": str(e)},
                "grade": None,
                "comment": "Vision AI call failed",
                "student_id": None,
                "outsider_student_id": None
            }
            if is_outsider:
                error_payload["outsider_student_id"] = entity_id
            else:
                error_payload["student_id"] = entity_id

            for q_id in all_question_ids:
                self.db.create_ai_model_run(**{**error_payload, "question_id": q_id})
            return {'success': False, 'tokens': {'prompt_tokens': 0, 'completion_tokens': 0, 'total_tokens': 0}}

    async def _grade_entire_submission_for_entity(
        self, job_id: str, entity_id: str, is_outsider: bool, answer_sheet_path: str, content_type: str, config: Union[assessment_model.AssessmentConfig, assessment_model.AssessmentConfigV2], user_id: str
    ) -> Dict:
        all_questions = [q for s in config.sections for q in s.questions] if isinstance(config, assessment_model.AssessmentConfigV2) else config.questions
        try:
            # Read the answer sheet file as bytes for vision processing
            with open(answer_sheet_path, "rb") as f:
                file_bytes = f.read()

            questions_json_str = json.dumps([q.model_dump(by_alias=True) for q in all_questions], indent=2)

            print(f"[GRADE-VISION] job={job_id} entity={entity_id} is_outsider={is_outsider} file_size={len(file_bytes)} questions={len(all_questions)}")

            answer_key_context = "The correct answers are included within each question object in the JSON below. Use these as the ground truth."
            if isinstance(config, assessment_model.AssessmentConfigV2) and config.gradingMode != assessment_model.GradingMode.ANSWER_KEY_PROVIDED:
                 answer_key_context = "No answer key was provided. Grade based on general knowledge."

            prompt = prompt_library.STUDENT_CENTRIC_GRADING_PROMPT.format(questions_json=questions_json_str, answer_key_context=answer_key_context)

            grading_tasks = [self._invoke_single_grading_run_vision(prompt, file_bytes, content_type, job_id, entity_id, is_outsider, i, user_id) for i in range(3)]
            results = await asyncio.gather(*grading_tasks)

            # Accumulate tokens from all 3 runs
            student_total_tokens = {'prompt_tokens': 0, 'completion_tokens': 0, 'total_tokens': 0}
            for result in results:
                if result.get('tokens'):
                    student_total_tokens['prompt_tokens'] += result['tokens']['prompt_tokens']
                    student_total_tokens['completion_tokens'] += result['tokens']['completion_tokens']
                    student_total_tokens['total_tokens'] += result['tokens']['total_tokens']

            # Get student name for summary log
            student_name = "Unknown"
            if not is_outsider:
                student_obj = self.db.get_student_by_id(entity_id, user_id)
                if student_obj:
                    student_name = student_obj.name
            else:
                outsider_obj = self.db.get_outsider_student_by_id(entity_id, user_id)
                if outsider_obj:
                    student_name = outsider_obj.name

            print(f"[TOKEN-USAGE] STUDENT-TOTAL [{student_name}] - Prompt: {student_total_tokens['prompt_tokens']}, Completion: {student_total_tokens['completion_tokens']}, Total: {student_total_tokens['total_tokens']}")

            # Return token usage for assessment-level aggregation
            entity_tokens = student_total_tokens.copy()

            for question in all_questions:
                runs = self.db.get_ai_model_runs_for_question(job_id, entity_id, question.id, is_outsider)

                grades = [run.grade for run in runs]
                comments = [run.comment for run in runs]
                max_score = question.maxScore if question.maxScore else 10.0

                decimal_grades = [Decimal(str(g)) if g is not None else None for g in grades]
                consensus_result = finalize_question(decimal_grades, comments, max_score)

                print(f"[CONSENSUS] job={job_id} entity={entity_id} q={question.id} grades={[str(g) for g in grades]} status={consensus_result['status']} final={consensus_result.get('grade')} by={consensus_result.get('finalized_by')}")

                if is_outsider:
                    self.db.update_outsider_result_grade(
                        job_id=job_id, outsider_student_id=entity_id, question_id=question.id,
                        grade=consensus_result.get('grade'),
                        feedback=consensus_result.get('feedback'),
                        status=consensus_result['status'].value,
                        finalized_by=consensus_result.get('finalized_by'),
                        user_id=user_id
                    )
                else:
                    self.db.update_student_result_with_grade(
                        job_id=job_id, student_id=entity_id, question_id=question.id,
                        grade=consensus_result.get('grade'),
                        feedback=consensus_result.get('feedback'),
                        status=consensus_result['status'].value,
                        finalized_by=consensus_result.get('finalized_by'),
                        user_id=user_id
                    )

            return entity_tokens

        except Exception as e:
            print(f"CRITICAL ERROR processing submission for entity {entity_id} in job {job_id}: {e}")
            # for q in all_questions:
                # self.db.update_result_status(job_id, entity_id, q.id, "FAILED", user_id)
            return {'prompt_tokens': 0, 'completion_tokens': 0, 'total_tokens': 0}

    async def _grade_entity_with_semaphore(self, semaphore: asyncio.Semaphore, job_id: str, entity: Dict, config: Union[assessment_model.AssessmentConfig, assessment_model.AssessmentConfigV2], user_id: str) -> Dict:
        """A wrapper to acquire the semaphore before grading an entity."""
        async with semaphore:
            await asyncio.sleep(1) # Stagger the start of API calls
            tokens = await self._grade_entire_submission_for_entity(
                job_id,
                entity['entity_id'],
                entity['is_outsider'],
                entity['answer_sheet_path'],
                entity['content_type'],
                config,
                user_id
            )
            return tokens

    async def process_assessment_job(self, job_id: str, user_id: str):
        semaphore = asyncio.Semaphore(2)
        assessment_total_tokens = {'prompt_tokens': 0, 'completion_tokens': 0, 'total_tokens': 0}
        try:
            self.db.update_job_status(job_id, user_id, assessment_model.JobStatus.PROCESSING.value)

            # Check if this is a manual upload job - if so, skip file matching as it's already done
            job_record = self.db.get_assessment_job(job_id, user_id)
            config_data = job_record.config if isinstance(job_record.config, dict) else {}
            is_manual_upload = config_data.get('is_manual_upload', False)

            if not is_manual_upload:
                # Only match files for batch uploads
                await analytics_and_matching.match_files_to_students(self.db, job_id, user_id)
            else:
                print(f"[MANUAL-UPLOAD-PROCESSING] Skipping file matching for manual upload job {job_id}")

            job_record = self.db.get_assessment_job(job_id, user_id)
            if not job_record:
                raise ValueError(f"Job {job_id} not found for user {user_id} during processing.")

            config = analytics_and_matching.get_validated_config_from_job(job_record)
            entities_to_grade = self.db.get_entities_with_paths(job_id, user_id)

            print(f"[TOKEN-USAGE] ========== ASSESSMENT START: {config.assessmentName} (Job: {job_id}) ==========")

            grading_tasks = [
                self._grade_entity_with_semaphore(
                    semaphore, job_id, entity, config, user_id
                ) for entity in entities_to_grade
            ]
            grading_results = await asyncio.gather(*grading_tasks)

            # Accumulate tokens from all students
            for entity_tokens in grading_results:
                if entity_tokens:
                    assessment_total_tokens['prompt_tokens'] += entity_tokens.get('prompt_tokens', 0)
                    assessment_total_tokens['completion_tokens'] += entity_tokens.get('completion_tokens', 0)
                    assessment_total_tokens['total_tokens'] += entity_tokens.get('total_tokens', 0)

            print(f"[TOKEN-USAGE] ========== ASSESSMENT GRADING COMPLETE ==========")
            print(f"[TOKEN-USAGE] Total Students Graded: {len(entities_to_grade)}")
            print(f"[TOKEN-USAGE] ASSESSMENT-TOTAL - Prompt: {assessment_total_tokens['prompt_tokens']}, Completion: {assessment_total_tokens['completion_tokens']}, Total: {assessment_total_tokens['total_tokens']}")
            print(f"[TOKEN-USAGE] ========== END ASSESSMENT: {config.assessmentName} ==========")
            
            if self.db.are_any_questions_pending_review(job_id, user_id):
                self.db.update_job_status(job_id, user_id, assessment_model.JobStatus.PENDING_REVIEW.value)
            else:
                self.db.update_job_status(job_id, user_id, assessment_model.JobStatus.SUMMARIZING.value)
                await self._generate_analytic_summary(job_id, user_id)
                self.db.update_job_status(job_id, user_id, assessment_model.JobStatus.COMPLETED.value)
        except Exception as e:
            print(f"CRITICAL ERROR processing job {job_id} for user {user_id}: {e}")
            self.db.update_job_status(job_id, user_id, assessment_model.JobStatus.FAILED.value)

    def get_all_assessment_jobs_summary(self, user_id: str) -> Dict:
        all_jobs = self.db.get_all_assessment_jobs(user_id=user_id)
        all_results = self.db.get_all_results_for_user(user_id=user_id)
        all_classes_objs = self.db.get_all_classes(user_id=user_id)
        all_classes_map = {c.id: c.name for c in all_classes_objs}
        summaries = data_assembly._assemble_job_summaries(all_jobs, all_results, all_classes_map)
        return {"assessments": summaries}

    def get_full_job_results(self, job_id: str, user_id: str) -> Optional[Dict]:
        job_record = self.db.get_assessment_job(job_id=job_id, user_id=user_id)
        if not job_record:
            return None
        config_v2 = analytics_and_matching.normalize_config_to_v2(job_record)
        class_students = self.db.get_students_by_class_id(class_id=config_v2.classId, user_id=user_id)
        all_results_for_job = self.db.get_all_results_for_job(job_id=job_id, user_id=user_id)
        final_results_dict = data_assembly._build_results_dictionary(class_students, config_v2, all_results_for_job)
        students_list = [{"id": s.id, "name": s.name, "answerSheetPath": self.db.get_student_result_path(job_id, s.id, user_id)} for s in class_students]
        analytics_data = None
        if job_record.status == assessment_model.JobStatus.COMPLETED.value:
            analytics_data = analytics_and_matching.calculate_analytics(all_results_for_job, config_v2)
        return {
            "jobId": job_record.id, "assessmentName": config_v2.assessmentName,
            "status": job_record.status, "config": config_v2,
            "students": students_list, "results": final_results_dict,
            "analytics": analytics_data, "aiSummary": job_record.ai_summary
        }

    async def _generate_analytic_summary(self, job_id: str, user_id: str):
        job_results = self.get_full_job_results(job_id=job_id, user_id=user_id)
        if not job_results or not job_results.get('analytics'):
            return
        prompt = prompt_library.ANALYTICS_SUMMARY_PROMPT.format(analytics_json=json.dumps(job_results['analytics'], indent=2))
        summary_text = await gemini_service.generate_text(prompt, temperature=0.6)
        self.db.update_job_with_summary(job_id=job_id, user_id=user_id, summary=summary_text)

    def get_assessment_results_overview(self, job_id: str, user_id: str) -> assessment_model.AssessmentResultsOverviewResponse:
        job = self.db.get_assessment_job(job_id=job_id, user_id=user_id)
        if not job:
            raise ValueError(f"Job with ID {job_id} not found or access denied.")
        config = analytics_and_matching.normalize_config_to_v2(job)
        students = self.db.get_students_by_class_id(config.classId, user_id)
        results = self.db.get_all_results_for_job(job_id, user_id)
        results_by_student = {}
        for r in results:
            if r.student_id not in results_by_student:
                results_by_student[r.student_id] = []
            results_by_student[r.student_id].append(r)
        students_ai_graded = []
        students_pending = []
        for student in students:
            student_results = results_by_student.get(student.id, [])
            pending_count = sum(1 for r in student_results if r.status == ResultStatus.PENDING_REVIEW.value)
            if pending_count > 0:
                students_pending.append(assessment_model.StudentPendingSummary(
                    student_id=student.id, name=student.name, num_pending=pending_count
                ))
            else:
                total_score = sum(r.grade for r in student_results if r.grade is not None)
                students_ai_graded.append(assessment_model.StudentAIGradedSummary(
                    student_id=student.id, name=student.name, total_score=total_score
                ))
        return assessment_model.AssessmentResultsOverviewResponse(
            job_id=job.id, assessment_name=config.assessmentName, status=job.status,
            students_ai_graded=students_ai_graded, students_pending=students_pending
        )

    def get_student_assessment_for_review(self, job_id: str, entity_id: str, user_id: str) -> assessment_model.StudentReviewResponse:
        job = self.db.get_assessment_job(job_id=job_id, user_id=user_id)
        if not job:
            raise ValueError(f"Job with ID {job_id} not found or access denied.")

        # Determine if the entity is a rostered student or an outsider
        student_name = None
        student_db_id = None
        outsider_db_id = None

        rostered_student = self.db.get_student_by_id(entity_id, user_id)
        if rostered_student:
            student_name = rostered_student.name
            student_db_id = rostered_student.id
        else:
            outsider_student = self.db.get_outsider_student_by_id(entity_id, user_id)
            if not outsider_student:
                raise ValueError(f"Entity with ID {entity_id} not found.")
            student_name = outsider_student.name
            outsider_db_id = outsider_student.id

        config = analytics_and_matching.normalize_config_to_v2(job)
        all_results_for_job = self.db.get_all_results_for_job(job_id, user_id)

        # Filter results based on whether it's a rostered or outsider student
        if student_db_id:
            entity_results = [r for r in all_results_for_job if r.student_id == student_db_id]
        else:
            entity_results = [r for r in all_results_for_job if r.outsider_student_id == outsider_db_id]

        questions_for_review = []
        all_questions = [q for s in config.sections for q in s.questions]
        for question in all_questions:
            result = next((r for r in entity_results if r.question_id == question.id), None)
            if result:
                questions_for_review.append(assessment_model.QuestionForReview(
                    question_id=question.id, question_text=question.text,
                    max_score=question.maxScore or 0, student_answer=result.extractedAnswer,
                    status=result.status, grade=result.grade, feedback=result.feedback
                ))

        return assessment_model.StudentReviewResponse(
            job_id=job.id, student_id=entity_id, student_name=student_name,
            assessment_name=config.assessmentName, config=config, per_question=questions_for_review
        )

    async def get_combined_overview(self, job_id: str, user_id: str) -> List[assessment_model.StudentResultRow]:
        job = self.db.get_assessment_job(job_id=job_id, user_id=user_id)
        if not job:
            raise ValueError(f"Job with ID {job_id} not found or access denied.")

        config = analytics_and_matching.normalize_config_to_v2(job)

        # Calculate the total possible score from the config
        max_total_score = sum(q.maxScore for section in config.sections for q in section.questions if q.maxScore is not None)

        # 1. Fetch all data sources
        roster = self.db.get_students_by_class_id(config.classId, user_id)
        outsiders = self.db.get_all_outsider_students_for_job(job_id, user_id)
        all_results = self.db.get_all_results_for_job(job_id, user_id)

        # 2. Group results by entity ID
        results_by_roster_id: Dict[str, List] = {s.id: [] for s in roster}
        results_by_outsider_id: Dict[str, List] = {o.id: [] for o in outsiders}
        for r in all_results:
            if r.student_id and r.student_id in results_by_roster_id:
                results_by_roster_id[r.student_id].append(r)
            elif r.outsider_student_id and r.outsider_student_id in results_by_outsider_id:
                results_by_outsider_id[r.outsider_student_id].append(r)

        rows: List[assessment_model.StudentResultRow] = []

        # 3. Process rostered students
        for student in roster:
            student_results = results_by_roster_id.get(student.id, [])
            if not student_results:
                rows.append(assessment_model.StudentResultRow(
                    entity_id=student.id, student_id=student.studentId, student_name=student.name,
                    status="ABSENT", is_absent=True, is_outsider=False, max_total_score=max_total_score
                ))
                continue

            statuses = {res.status for res in student_results}
            any_pending = ResultStatus.PENDING_REVIEW.value in statuses
            any_teacher = any(res.finalized_by == FinalizedBy.TEACHER.value for res in student_results)

            status = "PENDING_REVIEW" if any_pending else ("TEACHER_GRADED" if any_teacher else "AI_GRADED")
            total_score = sum(float(res.grade) for res in student_results if res.grade is not None) if not any_pending else None

            rows.append(assessment_model.StudentResultRow(
                entity_id=student.id, student_id=student.studentId, student_name=student.name,
                status=status, total_score=total_score, is_outsider=False, max_total_score=max_total_score
            ))

        # 4. Process outsider students
        for outsider in outsiders:
            outsider_results = results_by_outsider_id.get(outsider.id, [])
            if not outsider_results:
                continue # Should not happen, as outsiders are created with results

            statuses = {res.status for res in outsider_results}
            any_pending = ResultStatus.PENDING_REVIEW.value in statuses
            any_teacher = any(res.finalized_by == FinalizedBy.TEACHER.value for res in outsider_results)

            status = "PENDING_REVIEW" if any_pending else ("TEACHER_GRADED" if any_teacher else "AI_GRADED")
            total_score = sum(float(res.grade) for res in outsider_results if res.grade is not None) if not any_pending else None

            rows.append(assessment_model.StudentResultRow(
                entity_id=outsider.id,
                student_id="Outsider", # Keep this for display grouping
                student_name=outsider.name, # Use the actual name from the DB
                status=status,
                total_score=total_score,
                is_outsider=True,
                max_total_score=max_total_score
            ))

        return rows

    async def apply_teacher_edit(self, job_id: str, entity_id: str, question_id: str, grade: float, feedback: str, user_id: str) -> Dict:
        job = self.db.get_assessment_job(job_id=job_id, user_id=user_id)
        if not job:
            raise ValueError(f"Job with ID {job_id} not found or access denied.")
        config = analytics_and_matching.normalize_config_to_v2(job)

        # Validate question and grade
        all_questions = [q for s in config.sections for q in s.questions]
        question_config = next((q for q in all_questions if q.id == question_id), None)
        if not question_config:
            raise ValueError(f"Question with ID {question_id} not found in assessment config.")
        max_score = question_config.maxScore if question_config.maxScore is not None else 100.0
        if grade < 0 or grade > max_score:
            raise ValueError(f"Grade must be between 0 and {max_score}")

        # Determine if entity is rostered or outsider
        rostered_student = self.db.get_student_by_id(entity_id, user_id)
        outsider_student = None if rostered_student else self.db.get_outsider_student_by_id(entity_id, user_id)

        if not rostered_student and not outsider_student:
            raise ValueError(f"Entity {entity_id} not found or not part of this job.")

        if rostered_student:
            self.db.update_student_result_with_grade(
                job_id=job_id, student_id=entity_id, question_id=question_id, grade=grade,
                feedback=feedback, status=ResultStatus.TEACHER_GRADED.value,
                finalized_by=FinalizedBy.TEACHER.value, user_id=user_id
            )
            display_id = rostered_student.studentId
        else: # Is outsider
            self.db.update_outsider_result_grade(
                job_id=job_id, outsider_student_id=entity_id, question_id=question_id, grade=grade,
                feedback=feedback, status=ResultStatus.TEACHER_GRADED.value,
                finalized_by=FinalizedBy.TEACHER.value, user_id=user_id
            )
            display_id = "Outsider"

        # After saving, check if the entire job is now complete
        if not self.db.are_any_questions_pending_review(job_id, user_id):
            self.db.update_job_status(job_id, user_id, assessment_model.JobStatus.COMPLETED.value)
            print(f"Job {job_id} status updated to COMPLETED as all reviews are done.")

        # Calculate the updated total score for the entity and return
        all_results_for_job = self.db.get_all_results_for_job(job_id, user_id)
        if rostered_student:
            entity_results = [r for r in all_results_for_job if r.student_id == entity_id]
        else:
            entity_results = [r for r in all_results_for_job if r.outsider_student_id == entity_id]

        pending_left = any(r.status == ResultStatus.PENDING_REVIEW.value for r in entity_results)
        total_score = sum(float(r.grade) for r in entity_results if r.grade is not None)

        return {"studentId": display_id, "pendingLeft": pending_left, "totalScore": total_score}

    async def build_student_report_docx(self, job_id: str, entity_id: str, user_id: str) -> BytesIO:
        job = self.db.get_assessment_job(job_id=job_id, user_id=user_id)
        if not job:
            raise ValueError(f"Job with ID {job_id} not found or access denied.")

        config = analytics_and_matching.normalize_config_to_v2(job)

        # Determine if the entity is a rostered student or an outsider
        entity_name = None
        display_id = None
        db_id = entity_id
        is_outsider = False

        rostered_student = self.db.get_student_by_id(entity_id, user_id)
        if rostered_student:
            entity_name = rostered_student.name
            display_id = rostered_student.studentId
        else:
            outsider_student = self.db.get_outsider_student_by_id(entity_id, user_id)
            if not outsider_student:
                raise ValueError(f"Entity with ID {entity_id} not found.")
            entity_name = outsider_student.name
            display_id = "Outsider"
            is_outsider = True

        all_results_for_job = self.db.get_all_results_for_job(job_id, user_id)
        if is_outsider:
            entity_results = [r for r in all_results_for_job if r.outsider_student_id == db_id]
        else:
            entity_results = [r for r in all_results_for_job if r.student_id == db_id]

        results_map = {r.question_id: r for r in entity_results}
        all_questions = [q for s in config.sections for q in s.questions]

        doc = Document()
        doc.add_heading(f"Assessment Report: {config.assessmentName}", level=1)
        doc.add_paragraph(f"Student: {entity_name} (ID: {display_id})")
        doc.add_paragraph(f"Class: {self.db.get_class_by_id(config.classId, user_id).name}\n")

        total_score = 0.0
        max_total_score = 0.0

        for i, question in enumerate(all_questions):
            doc.add_heading(f"Question {i+1}: {question.text}", level=2)
            max_score = question.maxScore or 0
            max_total_score += max_score
            doc.add_paragraph(f"Max Score: {max_score}")

            result = results_map.get(question.id)

            p_answer = doc.add_paragraph()
            p_answer.add_run("Student's Answer:\n").bold = True
            student_answer = result.extractedAnswer if result and result.extractedAnswer else "Not answered"
            p_answer.add_run(student_answer)

            p_grade = doc.add_paragraph()
            p_grade.add_run("Grade: ").bold = True
            if result and result.grade is not None:
                p_grade.add_run(f"{result.grade} / {max_score}")
                total_score += float(result.grade)
            else:
                p_grade.add_run("Pending Review")

            p_feedback = doc.add_paragraph()
            p_feedback.add_run("Feedback:\n").bold = True
            feedback = result.feedback if result and result.feedback else "No feedback provided."
            p_feedback.add_run(feedback)
            doc.add_paragraph()

        doc.add_heading(f"Final Score: {total_score} / {max_total_score}", level=2)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    async def distribute_scores_with_ai(self, config: assessment_model.AssessmentConfigV2, total_marks: int) -> assessment_model.AssessmentConfigV2:
        """
        Uses AI to distribute a total score across all questions in a configuration.
        """
        try:
            config_json_str = config.model_dump_json(by_alias=True)

            prompt = prompt_library.AI_SCORE_DISTRIBUTION_PROMPT.format(
                total_marks=total_marks,
                questions_json=config_json_str
            )

            # Use the JSON-specific generator function for a reliable response
            response_json = await gemini_service.generate_json(prompt, temperature=0.3)

            # Validate and return the updated config
            updated_config = assessment_model.AssessmentConfigV2.model_validate(response_json)
            return updated_config

        except (ValueError, json.JSONDecodeError) as e:
            print(f"Error processing AI response for score distribution: {e}")
            # In case of error, return the original config to avoid breaking the flow
            return config
        except Exception as e:
            print(f"An unexpected error occurred during AI score distribution: {e}")
            return config

def get_assessment_service(db: DatabaseService = Depends(get_db_service)):
    return AssessmentService(db=db)